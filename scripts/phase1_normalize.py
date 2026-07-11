#!/usr/bin/env python3
"""
Phase 1 — Source normalization.

Reads every .docx in sources/notebooklm-exports/raw_extracted/, extracts a clean
Markdown rendering (headings + paragraphs + tables) into
content/normalized/normalized_sources/, and builds a machine-readable
source_index.json with:
  - stable source id, original filename, normalized path
  - size, word count, sha256 of original + normalized text
  - classification: approved_content | excluded_artifact
  - exact-duplicate and near-duplicate detection
  - lightweight blueprint-section hints (keyword signal, used later in Phase 3)

Non-.docx artifacts (the stray large .pptx) are recorded and EXCLUDED from the
approved-generation layer.
"""
from __future__ import annotations
import hashlib
import json
import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document  # python-docx

ROOT = Path("/Users/omarbaba/Library/CloudStorage/OneDrive-Personal/tala blood bank")
RAW = ROOT / "sources" / "notebooklm-exports" / "raw_extracted"
OUT_DIR = ROOT / "content" / "normalized" / "normalized_sources"
INDEX = ROOT / "content" / "normalized" / "source_index.json"

OUT_DIR.mkdir(parents=True, exist_ok=True)

# Blueprint section keyword signals (for a coarse Phase-3 pre-map, not authoritative)
SECTION_KEYWORDS = {
    "1. Clinical Practice": ["autoimmune hemolytic", "waiha", "cold agglutinin", "pnh", "immune thrombocytopenia", "itp", "hit", "heparin-induced", "dic", "coagulation", "warfarin", "von willebrand", "granulocyte"],
    "2. Cell and Tissue Therapy": ["hla", "histocompatibility", "tissue banking", "adoptive immunotherapy", "car-t", "gene therapy", "regenerative"],
    "3. RBCs and RBC Components": ["erythropoietin", "erythropoiesis", "red cell metabolism", "additive solution", "abo typing", "rh typing", "antiglobulin", "dat", "abh", "lewis", "kell", "duffy", "kidd", "mns", "antibody identification", "crossmatch", "blood group"],
    "4. Anemia and Red Blood Cell Transfusion": ["transfusion threshold", "restrictive transfusion", "anemia", "hemoglobin trigger", "decision making"],
    "5. Apheresis": ["apheresis", "plasma exchange", "tpe", "red cell exchange", "photopheresis", "ecp", "phlebotomy", "leukapheresis"],
    "6. Hazards of Transfusion": ["hemolytic transfusion reaction", "febrile", "fnhtr", "allergic", "taco", "trali", "ta-gvhd", "anaphylactic", "iron overload", "massive transfusion", "citrate", "dyspnea", "tad"],
    "7. Plasma Components and Derivatives": ["fresh frozen plasma", "ffp", "cryoprecipitate", "albumin", "pcc", "prothrombin complex", "ivig", "rh immune globulin", "rhig", "fibrinogen"],
    "8. Infectious Hazards of Transfusion": ["hepatitis", "hiv", "htlv", "cmv", "west nile", "parvovirus", "zika", "chagas", "babesiosis", "malaria", "bacterial contamination", "pathogen inactivation", "prion", "syphilis"],
    "9. Blood Donors and Blood Collection": ["donor recruitment", "donor eligibility", "deferral", "donor reaction", "component separation", "leukoreduction", "collection process", "phlebotomy-related"],
    "10. Surgery Patients": ["cell salvage", "normovolemic hemodilution", "perioperative", "trauma", "massive transfusion protocol", "burn", "solid organ transplant"],
    "11. Biovigilance": ["biovigilance", "hemovigilance", "trim", "immunomodulation", "traceability"],
    "12. Platelets": ["thrombopoiesis", "megakaryocyte", "platelet transfusion", "platelet antigen", "hpa", "platelet refractoriness", "platelet storage"],
    "13. Neutrophils": ["neutrophil", "granulocyte transfusion", "myelopoietic"],
    "14. Intravascular Cell Kinetics": ["cell kinetics", "post-transfusion recovery", "labeled cells", "chromium"],
    "15. Obstetric and Pediatric Patients": ["hdfn", "hemolytic disease", "rhig", "intrauterine transfusion", "neonatal", "prematurity", "sickle cell", "thalassemia", "hemophilia", "exchange transfusion"],
    "16. HPC Transplantation": ["hpc", "hematopoietic progenitor", "stem cell", "mobilization", "plerixafor", "cord blood", "engraftment", "cd34", "cryopreservation", "dmso", "passenger lymphocyte"],
    "17. Administration and Laboratory Management": ["quality management", "cgmp", "fda", "aabb", "accreditation", "transfusion committee", "medical director", "regulatory", "informed consent", "hipaa", "sop", "error management"],
}


def sha256_bytes(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()


def slugify(name: str) -> str:
    s = re.sub(r"\.docx$", "", name, flags=re.IGNORECASE)
    s = s.replace("’", "'")
    s = re.sub(r"[^a-zA-Z0-9]+", "-", s).strip("-").lower()
    return s[:80]


def docx_to_markdown(path: Path) -> str:
    doc = Document(str(path))
    lines: list[str] = []

    def emit_para(p) -> None:
        text = p.text.strip()
        if not text:
            return
        style = (p.style.name or "").lower() if p.style else ""
        if style.startswith("title"):
            lines.append(f"# {text}")
        elif style.startswith("heading 1") or style == "heading":
            lines.append(f"## {text}")
        elif style.startswith("heading 2"):
            lines.append(f"### {text}")
        elif style.startswith("heading 3"):
            lines.append(f"#### {text}")
        elif style.startswith("heading"):
            lines.append(f"##### {text}")
        elif style.startswith("list") or p.text.lstrip().startswith(("•", "-", "*")):
            lines.append(f"- {text.lstrip('•-* ')}")
        else:
            lines.append(text)

    for p in doc.paragraphs:
        emit_para(p)

    # Tables (render as pipe tables)
    for ti, table in enumerate(doc.tables):
        lines.append("")
        lines.append(f"<!-- table {ti + 1} -->")
        for ri, row in enumerate(table.rows):
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
            if ri == 0:
                lines.append("| " + " | ".join(["---"] * len(cells)) + " |")

    return "\n\n".join(lines).strip() + "\n"


def word_shingles(text: str, n: int = 3) -> set:
    words = re.findall(r"[a-z0-9]+", text.lower())
    return {" ".join(words[i:i + n]) for i in range(max(0, len(words) - n + 1))}


def jaccard(a: set, b: set) -> float:
    if not a or not b:
        return 0.0
    inter = len(a & b)
    union = len(a | b)
    return inter / union if union else 0.0


def section_hints(text: str) -> list[str]:
    low = text.lower()
    scored = []
    for section, kws in SECTION_KEYWORDS.items():
        hits = sum(1 for k in kws if k in low)
        if hits:
            scored.append((section, hits))
    scored.sort(key=lambda x: -x[1])
    return [f"{s} ({h})" for s, h in scored[:5]]


def main() -> None:
    docx_files = sorted(RAW.glob("*.docx"))
    other_files = [p for p in RAW.iterdir() if p.is_file() and p.suffix.lower() != ".docx"]

    records = []
    shingle_map = {}

    for i, path in enumerate(docx_files, start=1):
        raw_bytes = path.read_bytes()
        try:
            md = docx_to_markdown(path)
        except Exception as e:  # noqa: BLE001
            md = f"<!-- extraction error: {e} -->\n"
        slug = slugify(path.name)
        out_path = OUT_DIR / f"{slug}.md"
        # avoid slug collisions
        if out_path.exists() and out_path.read_text() != md:
            out_path = OUT_DIR / f"{slug}-{i:02d}.md"
        out_path.write_text(md)

        words = re.findall(r"[a-z0-9]+", md.lower())
        text_hash = sha256_bytes(md.encode())
        rec = {
            "id": f"SRC-{i:03d}",
            "original_filename": path.name,
            "normalized_path": str(out_path.relative_to(ROOT)),
            "classification": "approved_content",
            "docx_sha256": sha256_bytes(raw_bytes),
            "normalized_text_sha256": text_hash,
            "size_bytes": len(raw_bytes),
            "word_count": len(words),
            "section_hints": section_hints(md),
            "exact_duplicate_of": None,
            "near_duplicates": [],
        }
        records.append(rec)
        shingle_map[rec["id"]] = (word_shingles(md), text_hash)

    # Exact duplicate detection (identical normalized text)
    by_texthash: dict[str, str] = {}
    for rec in records:
        th = rec["normalized_text_sha256"]
        if th in by_texthash:
            rec["exact_duplicate_of"] = by_texthash[th]
        else:
            by_texthash[th] = rec["id"]

    # Near-duplicate detection (Jaccard on 3-gram shingles)
    ids = [r["id"] for r in records]
    NEAR = 0.80
    for a_idx in range(len(ids)):
        for b_idx in range(a_idx + 1, len(ids)):
            ra, rb = records[a_idx], records[b_idx]
            if ra["exact_duplicate_of"] or rb["exact_duplicate_of"]:
                continue
            sim = jaccard(shingle_map[ra["id"]][0], shingle_map[rb["id"]][0])
            if sim >= NEAR:
                ra["near_duplicates"].append({"id": rb["id"], "jaccard": round(sim, 3)})
                rb["near_duplicates"].append({"id": ra["id"], "jaccard": round(sim, 3)})

    # Excluded artifacts (non-docx)
    excluded = []
    for p in other_files:
        excluded.append({
            "id": f"EXC-{p.suffix.lstrip('.').upper()}",
            "original_filename": p.name,
            "classification": "excluded_artifact",
            "reason": "Non-educational artifact (not part of the BB/TM study corpus); "
                      "unrelated presentation file swept into the Drive export.",
            "size_bytes": p.stat().st_size,
            "docx_sha256": None,
        })

    n_exact = sum(1 for r in records if r["exact_duplicate_of"])
    n_near_pairs = sum(len(r["near_duplicates"]) for r in records) // 2
    unique_content = [r for r in records if not r["exact_duplicate_of"]]

    index = {
        "phase": 1,
        "generated_at_iso": datetime.now(tz=timezone.utc).isoformat(),
        "corpus": "NotebookLM BB/TM board-prep export",
        "counts": {
            "approved_content_docs": len(records),
            "unique_content_docs": len(unique_content),
            "exact_duplicates": n_exact,
            "near_duplicate_pairs": n_near_pairs,
            "excluded_artifacts": len(excluded),
            "total_words_unique": sum(r["word_count"] for r in unique_content),
        },
        "approved_sources": records,
        "excluded_artifacts": excluded,
    }
    INDEX.write_text(json.dumps(index, indent=2))
    print(f"Wrote {INDEX.relative_to(ROOT)}")
    print(f"  approved docs: {len(records)} | unique: {len(unique_content)} | "
          f"exact dups: {n_exact} | near-dup pairs: {n_near_pairs} | excluded: {len(excluded)}")
    print(f"  total unique words: {index['counts']['total_words_unique']:,}")


if __name__ == "__main__":
    main()
